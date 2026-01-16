import os
import re
import PyPDF2
import docx

from rag_summary import create_rag_summary

# =====================================================
# ENV CHECK
# =====================================================
IS_PRODUCTION = bool(os.environ.get("RENDER"))

if not IS_PRODUCTION:
    try:
        import ollama
        OLLAMA_AVAILABLE = True
    except ImportError:
        OLLAMA_AVAILABLE = False
else:
    OLLAMA_AVAILABLE = False


# =====================================================
# SAFETY HELPERS
# =====================================================
def contains_phone_or_email(text: str) -> bool:
    if re.search(r"\b\d{8,}\b", text):
        return True
    if re.search(r"\w+@\w+\.\w+", text):
        return True
    return False


# =====================================================
# RESUME TEXT EXTRACTION
# =====================================================
def get_text_from_resume(file_path):
    try:
        if file_path.endswith(".pdf"):
            text = ""
            with open(file_path, "rb") as f:
                pdf = PyPDF2.PdfReader(f)
                for page in pdf.pages:
                    if page.extract_text():
                        text += page.extract_text() + "\n"
            return text.strip() or None

        if file_path.endswith(".docx"):
            doc = docx.Document(file_path)
            parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text.strip())

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            parts.append(cell.text.strip())

            return "\n".join(parts) or None

    except Exception:
        return None


# =====================================================
# SANITIZER
# =====================================================
def sanitize_summary(summary, resume_text):
    bad_phrases = [
        "experienced professional with relevant skills",
        "skilled professional",
        "motivated professional",
    ]

    if not summary or any(b in summary.lower() for b in bad_phrases):
        return create_fallback_summary(resume_text)

    return summary


# =====================================================
# MAIN SUMMARY GENERATOR
# =====================================================
def create_resume_summary(text):
    if not text or len(text.strip()) < 50:
        return "• Resume content too brief for summary generation."

    if not OLLAMA_AVAILABLE:
        return create_fallback_summary(text)

    try:
        prompt = f"""
You are a professional resume writer with real recruiting experience.

Generate a PROFESSIONAL RESUME SUMMARY strictly from the resume below.

RULES:
- Bullet points only
- Minimum 5 bullets
- At least:
  • 2 project bullets
  • 2 skill bullets
  • 1 experience / responsibility bullet
- No generic phrases
- No contact info
- Preserve skill ratings if present

NEVER output:
"Experienced professional with relevant skills and experience"

RESUME TEXT:
{text[:1500]}
"""

        response = ollama.generate(
            model="llama3.2:1b",
            prompt=prompt,
            options={
                "num_predict": 300,
                "temperature": 0.1,
                "top_p": 0.8,
                "repeat_penalty": 1.1,
            },
        )

        summary = response.get("response", "").strip()
        return sanitize_summary(summary, text)

    except Exception as e:
        print("DEBUG: AI generation failed:", e)
        return create_fallback_summary(text)


# =====================================================
# FALLBACK SUMMARY (SAFE)
# =====================================================
def create_fallback_summary(text):
    lines = text.split("\n")
    bullets = []

    if re.search(r"\d+\+?\s*years?", text, re.I):
        bullets.append("• Experience across multiple software development cycles and real-world implementations.")

    techs = ["Java", "Python", "Flask", "Spring", "SQL", "JavaScript"]
    used = [t for t in techs if t.lower() in text.lower()]
    if used:
        bullets.append(f"• Hands-on experience working with technologies such as {', '.join(used[:4])}.")

    project_lines = [l for l in lines if "project" in l.lower() or "application" in l.lower()]
    for pl in project_lines[:2]:
        bullets.append(f"• {pl.strip()}")

    if len(bullets) < 5:
        bullets.append("• Actively contributed to development tasks, debugging, and feature enhancements.")

    return "\n".join(bullets[:7])


# =====================================================
# JOB MATCHING (UNCHANGED)
# =====================================================
def match_resume_with_job(resume_summary, job_title, job_description):
    if IS_PRODUCTION:
        job_text = f"{job_title} {job_description}".lower()
        resume_text = resume_summary.lower()

        keywords = ["python", "java", "react", "sql", "aws", "docker", "git"]
        matched = [k for k in keywords if k in job_text and k in resume_text]

        percent = int((len(matched) / max(len(keywords), 1)) * 100)
        return f"Match: {percent}% | Skills: {', '.join(matched)}"

    try:
        res = ollama.generate(
            model="llama3",
            prompt=f"JOB: {job_title}\n{job_description}\nCANDIDATE:\n{resume_summary}",
        )
        return res["response"]
    except Exception:
        return "AI matching failed"
