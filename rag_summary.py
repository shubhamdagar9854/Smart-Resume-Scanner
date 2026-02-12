import json
import re
import os
import time
import PyPDF2
import google.generativeai as genai
from dotenv import load_dotenv
from database import get_enhanced_prompt, add_enhanced_prompt, get_ai_feedback_by_type

load_dotenv()

# Gemini API key
GEMINI_API_KEY = os.getenv('GEMINI_API_KEY', 'your-api-key-here')

# Gemini client initialize (new syntax)
genai.configure(api_key=GEMINI_API_KEY)

# Model name
MODEL_NAME = 'models/gemini-2.5-flash'

# 1. Model Load karne ka function
def load_mistral_model():
    print("✅ Gemini API ready!")
    print("🌐 Using Google Gemini 2.0 Flash (Free tier)")
    return True

# 2. RAG-Enhanced AI Summary banane ka function
def generate_professional_summary_rag(resume_text: str) -> str:
    """Generate summary using RAG system with feedback integration"""
    
    # Get enhanced prompt from database
    enhanced_prompt = get_enhanced_prompt("summary")
    
    # Get relevant feedback
    feedback_list = get_ai_feedback_by_type("summary")
    feedback_context = ""
    if feedback_list:
        feedback_context = "\n\nIMPORTANT - Learn from these previous corrections:\n"
        for error, correction in feedback_list:
            feedback_context += f"Previous Error: {error}\nCorrection: {correction}\n\n"
    
    if enhanced_prompt:
        prompt = enhanced_prompt.format(resume_text=resume_text[:5000])
    else:
        # Original prompt if no enhanced version exists
        prompt = f"""
        Create EXACTLY 5 bullet point professional resume summary:
        
        Resume content: {resume_text[:5000]}
        
        Requirements:
        - Professional language
        - Key skills first
        - Experience summary  
        - Technical expertise
        - Soft skills last
        - Start each bullet with strong verb/action word
        
        Format exactly like this:
        * Skilled in [technologies]
        * Experienced [role] with expertise in  
        * Strong [skill] skills
        * Proficient in [methodologies/tools]  
        * Excellent [soft skill]
        
        Return ONLY the 5 bullet points, nothing else.
        """
    
    # Add feedback context to prompt
    if feedback_context:
        prompt += feedback_context
    
    # Retry logic for 503 errors
    max_retries = 3
    for attempt in range(max_retries):
        try:
            model = genai.GenerativeModel(MODEL_NAME)
            response = model.generate_content(prompt)
            return response.text
        except Exception as e:
            error_msg = str(e)
            
            # If model overloaded (503), wait and retry
            if '503' in error_msg or 'UNAVAILABLE' in error_msg:
                if attempt < max_retries - 1:
                    wait_time = (attempt + 1) * 2  # 2, 4, 6 seconds
                    print(f"Model busy, retrying in {wait_time}s... (Attempt {attempt + 1}/{max_retries})")
                    time.sleep(wait_time)
                    continue
            
            # If quota exceeded (429), use fallback summary
            if '429' in error_msg or 'RESOURCE_EXHAUSTED' in error_msg:
                return generate_fallback_summary(resume_text)
            
            # For other errors, return error message
            return f"Error generating summary: {error_msg}"
    
    # If all retries failed, use fallback
    return generate_fallback_summary(resume_text)


def generate_fallback_summary(resume_text: str) -> str:
    """Simple fallback summary when API fails"""
    lines = resume_text.split('\n')
    
    # Extract basic info
    skills = []
    experience = []
    
    for line in lines:
        line_lower = line.lower()
        if any(tech in line_lower for tech in ['java', 'python', 'sql', 'spring', 'docker', 'aws', 'kubernetes']):
            if len(line.strip()) < 100:
                skills.append(line.strip())
        if any(word in line_lower for word in ['lead', 'senior', 'consultant', 'engineer', 'developer']):
            if len(line.strip()) < 100:
                experience.append(line.strip())
    
    summary = "* " + " ".join(skills[:3]) if skills else "* Experienced professional with diverse technical skills"
    summary += "\n* " + (experience[0] if experience else "Software professional with proven track record")
    summary += "\n* Strong problem-solving and analytical capabilities"
    summary += "\n* Proficient in modern development tools and methodologies"
    summary += "\n* Excellent collaboration and communication skills"
    
    return summary

# 3. PDF se text nikalne ka function
def extract_text_from_resume(file_path):
    try:
        if not os.path.exists(file_path):
            return f"Error: File {file_path} nahi mili."
        
        # Try pdfplumber first (better for PDFs)
        if file_path.lower().endswith('.pdf'):
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    text = ""
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                    return text.strip() if text.strip() else "Error: No text extracted from PDF"
            except Exception as e:
                print(f"pdfplumber failed: {e}")
                # Fallback to PyPDF2
                try:
                    with open(file_path, 'rb') as file:
                        reader = PyPDF2.PdfReader(file)
                        text = "".join([page.extract_text() for page in reader.pages])
                        return text.strip() if text.strip() else "Error: No text extracted from PDF"
                except Exception as e2:
                    print(f"PyPDF2 also failed: {e2}")
                    return f"Error: Could not extract text from PDF: {e}"
        
        # Handle DOCX files
        elif file_path.lower().endswith('.docx'):
            try:
                import docx
                doc = docx.Document(file_path)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                return text.strip() if text.strip() else "Error: No text extracted from DOCX"
            except Exception as e:
                return f"Error: Could not extract text from DOCX: {e}"
        
        # Handle DOC files (if python-docx supports)
        elif file_path.lower().endswith('.doc'):
            try:
                import docx
                doc = docx.Document(file_path)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                return text.strip() if text.strip() else "Error: No text extracted from DOC"
            except Exception as e:
                return f"Error: Could not extract text from DOC: {e}"
        
        else:
            return f"Error: Unsupported file format: {file_path}"
            
    except Exception as e:
        return f"Error: {e}"

# 4. Main Process (Jo app.py call karega)
def main_process_resume(file_path):
    resume_text = extract_text_from_resume(file_path)
    summary = generate_professional_summary_rag(resume_text) 
    return {
        "ai_summary": summary, 
        "raw_text": resume_text
    }


def normalize_resume_json(raw_resume):
    """Clean and normalize resume JSON structure"""
    normalized = {
        "skills": raw_resume.get("skills", []),
        "projects": raw_resume.get("projects", []),
        "experience_years": raw_resume.get("experience_years", 0)
    }
    
    # Clean skills (remove garbage words)
    clean_skills = []
    garbage = {"any", "the", "and", "project", "skill", "work", "experience"}
    for skill in normalized["skills"]:
        skill_str = str(skill).strip().lower()
        if len(skill_str) > 2 and skill_str not in garbage:
            clean_skills.append(skill_str)
    
    normalized["skills"] = clean_skills
    return normalized


# 5. Matching Logic Functions
def get_ai_match_analysis(resume_text, job_text, match_percentage):
    """
    Get AI analysis for why candidate didn't get 100% match
    Uses AI to provide specific reasoning - ONLY AI, NO FALLBACK
    """
    try:
        if match_percentage >= 100:
            return "Perfect match! Candidate meets all job requirements."
        
        prompt = f"""
        Analyze why this candidate didn't get 100% match for the job:

        JOB DESCRIPTION:
        {job_text}

        RESUME:
        {resume_text}

        Current Match: {match_percentage}%

        Explain in 1-2 sentences why this candidate didn't achieve 100% match.
        Be specific about missing skills, experience, or qualifications.
        Focus on the most important gaps.

        Examples:
        - "Missing React framework experience which is specifically required."
        - "Lacks 5+ years of experience required for senior position."
        - "Missing database skills (SQL/NoSQL) mentioned in job requirements."
        """
        
        model = genai.GenerativeModel(MODEL_NAME)
        response = model.generate_content(prompt)
        
        if response and response.text:
            return response.text.strip()[:200]  # Limit to 200 characters
        else:
            return f"AI analysis unavailable. Current match: {match_percentage}%"
            
    except Exception as e:
        print(f"AI analysis failed: Quota exceeded")
        # Simple error message for quota exceeded
        if "429" in str(e) or "quota" in str(e).lower():
            return "AI quota exceeded. Please try again later."
        else:
            print(f"AI analysis failed: {e}")
            return f"AI analysis unavailable. Current match: {match_percentage}%"

def calculate_match_percentage_full_ai_rag(resume_text, job_text):
    """
    RAG-ENHANCED AI-BASED matching - Enhanced with proper feedback learning
    Uses Retrieval-Augmented Generation to improve AI performance from admin feedback
    """
    try:
        # Get relevant feedback for RAG
        feedback_list = get_ai_feedback_by_type("percentage")
        
        # Build enhanced prompt with feedback context
        if feedback_list:
            # Create RAG-enhanced prompt with feedback
            feedback_context = "\n\n=== LEARNING FROM PREVIOUS FEEDBACK ===\n"
            for i, (error, correction) in enumerate(feedback_list[:3]):  # Use last 3 feedbacks
                feedback_context += f"FEEDBACK {i+1}:\n"
                feedback_context += f"Previous Error: {error}\n"
                feedback_context += f"Admin Correction: {correction}\n"
                feedback_context += f"Learning: Apply this correction to current analysis\n\n"
            
            # Enhanced RAG prompt
            prompt = f"""
You are an intelligent resume matching AI that learns from feedback.

{feedback_context}

CURRENT TASK:
Calculate the match percentage between this RESUME and JOB DESCRIPTION:

===================
JOB DESCRIPTION:
{job_text}
===================

===================
RESUME:
{resume_text}
===================

CRITICAL INSTRUCTIONS:
1. Apply ALL learnings from the feedback above
2. Consider experience requirements seriously - if job requires 9+ years and candidate is student, DO NOT give high percentage
3. If feedback indicates 100% should be given for perfect skill matches, apply that logic
4. Analyze COMPLETE resume vs COMPLETE job description
5. Return ONLY a single number between 0-100

SCORING GUIDELINES:
- 90-100%: Perfect match - candidate exceeds ALL requirements including experience
- 75-89%: Strong match - candidate meets most requirements well
- 60-74%: Good match - candidate meets many requirements  
- 40-59%: Partial match - candidate meets some requirements
- 20-39%: Weak match - candidate meets few requirements
- 0-19%: No match - candidate doesn't meet requirements

Based on the feedback learnings above, calculate the FINAL match percentage.
Return ONLY the number:
"""
        else:
            # Fallback prompt if no feedback
            prompt = f"""
Calculate the match percentage between this RESUME and JOB DESCRIPTION:

===================
JOB DESCRIPTION:
{job_text}
===================

===================
RESUME:
{resume_text}
===================

Return ONLY a single number between 0-100 representing the match percentage.
"""
        
        print(f" RAG Enhanced Prompt: Using {len(feedback_list) if feedback_list else 0} feedback items")
        
        model = genai.GenerativeModel(MODEL_NAME)
        response = model.generate_content(prompt)
        
        # Extract percentage from AI response
        ai_response = response.text.strip()
        print(f" AI Raw Response: '{ai_response}'")
        
        # Find percentage in response - look for the LAST percentage number
        percentage_matches = re.findall(r'\d+(\.\d+)?', ai_response)
        print(f" Regex Matches Found: {percentage_matches}")
        
        if percentage_matches:
            # Take the LAST percentage (final answer after considering feedback)
            try:
                # Clean the percentage string - remove quotes and whitespace
                last_match = percentage_matches[-1].strip().strip('"').strip("'")
                print(f" Before Cleaning: '{percentage_matches[-1]}'")
                print(f" After Cleaning: '{last_match}'")
                percentage = float(last_match)
                # Ensure percentage is in valid range
                percentage = max(0, min(100, percentage))
                print(f" Final Percentage: {percentage}%")
                return round(percentage, 2)
            except (ValueError, IndexError) as e:
                print(f" Could not convert percentage to float: {percentage_matches[-1]} | Error: {e}")
                return 0.0
        else:
            print(" No percentage found in AI response")
            return 0.0
            
    except Exception as e:
        print(f"Full AI matching failed: {e}")
        # Fallback to basic analysis
        return 0.0

# 6. Self-Test Block
if __name__ == "__main__":
    # Testing mode
    test_pdf = "test_resume.pdf"
    if os.path.exists(test_pdf):
        print("Testing Mode...")
        load_mistral_model()
        res = main_process_resume(test_pdf)
        print("\n--- AI SUMMARY ---\n", res["ai_summary"])
        print("\n--- RAW TEXT (first 500 chars) ---\n", res["raw_text"][:500])
    else:
        print(f"  Test ke liye '{test_pdf}' file folder mein nahi hai.")
        print("Koi bhi resume PDF ko 'test_resume.pdf' naam se save karo aur phir se run karo!")